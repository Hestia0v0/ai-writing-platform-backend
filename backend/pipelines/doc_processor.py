"""
Document Processing Pipeline
=============================
Sequential stages executed by DocumentProcessor.process():

  1. Upload   – validate file type and size
  2. Parse    – extract raw text  (txt / pdf / docx)
  3. Clean    – normalise encoding, strip noise
  4. Chunk    – sliding-window splits ready for embedding / inference
  5. Score    – call ai_inference service; deterministic mock fallback
"""

from __future__ import annotations

import asyncio
import hashlib
import io
import logging
import re
import time
import unicodedata
import zipfile
from enum import Enum
from pathlib import Path
from typing import Protocol, runtime_checkable

import httpx
import jieba
from pydantic import BaseModel, Field

logger = logging.getLogger(__name__)

# Suppress jieba's initialization log output
jieba.setLogLevel(logging.WARNING)

_CJK_RANGES = (
    (0x4E00, 0x9FFF),   # CJK Unified Ideographs
    (0x3400, 0x4DBF),   # CJK Extension A
    (0x20000, 0x2A6DF), # CJK Extension B
    (0xF900, 0xFAFF),   # CJK Compatibility Ideographs
    (0x2E80, 0x2EFF),   # CJK Radicals Supplement
    (0x3040, 0x30FF),   # Hiragana + Katakana
    (0xAC00, 0xD7AF),   # Hangul syllables
)


def _is_cjk_char(ch: str) -> bool:
    cp = ord(ch)
    return any(lo <= cp <= hi for lo, hi in _CJK_RANGES)


def _is_cjk_text(text: str) -> bool:
    """Return True when more than 20% of non-whitespace characters are CJK."""
    non_ws = [c for c in text if not c.isspace()]
    if not non_ws:
        return False
    cjk_count = sum(1 for c in non_ws if _is_cjk_char(c))
    return cjk_count / len(non_ws) > 0.20


def _tokenize(text: str) -> list[str]:
    """Return a token list using jieba for CJK text, split() for Latin text."""
    if _is_cjk_text(text):
        return [t for t in jieba.cut(text) if t.strip()]
    return text.split()


# ── Constants ──────────────────────────────────────────────────────────────────

SUPPORTED_EXTENSIONS: frozenset[str] = frozenset({".txt", ".pdf", ".docx"})
DEFAULT_CHUNK_SIZE = 512     # words per chunk
DEFAULT_CHUNK_OVERLAP = 64   # word overlap between consecutive chunks
MAX_FILE_SIZE_MB = 20


# ── Stage Enum ─────────────────────────────────────────────────────────────────

class PipelineStage(str, Enum):
    UPLOAD = "upload"
    PARSE = "parse"
    CLEAN = "clean"
    CHUNK = "chunk"
    SCORE = "score"
    COMPLETE = "complete"


# ── Custom Exceptions ──────────────────────────────────────────────────────────

class PipelineError(Exception):
    """Base for all pipeline failures; carries the stage where it occurred."""

    def __init__(self, stage: PipelineStage, message: str) -> None:
        self.stage = stage
        super().__init__(f"[{stage.value}] {message}")


class UnsupportedFormatError(PipelineError):
    def __init__(self, extension: str) -> None:
        super().__init__(
            PipelineStage.UPLOAD,
            f"Unsupported format '{extension}'. Allowed: {sorted(SUPPORTED_EXTENSIONS)}",
        )


class MaliciousContentError(PipelineError):
    def __init__(self) -> None:
        super().__init__(PipelineStage.UPLOAD, "Malicious content detected in uploaded file.")


class FileSizeError(PipelineError):
    def __init__(self, size_mb: float) -> None:
        super().__init__(
            PipelineStage.UPLOAD,
            f"File size {size_mb:.1f} MB exceeds the {MAX_FILE_SIZE_MB} MB limit",
        )


class ParseError(PipelineError):
    def __init__(self, filename: str, detail: str) -> None:
        super().__init__(PipelineStage.PARSE, f"Cannot parse '{filename}': {detail}")


class InferenceError(PipelineError):
    def __init__(self, detail: str) -> None:
        super().__init__(PipelineStage.SCORE, f"AI inference failed: {detail}")


# ── Domain Models ──────────────────────────────────────────────────────────────

class UploadedDocument(BaseModel):
    model_config = {"arbitrary_types_allowed": True}

    document_id: str
    filename: str
    extension: str
    size_bytes: int
    content: bytes


class ParsedDocument(BaseModel):
    document_id: str
    filename: str
    raw_text: str
    page_count: int = 1


class CleanedDocument(BaseModel):
    document_id: str
    text: str
    char_count: int
    word_count: int


class TextChunk(BaseModel):
    chunk_index: int
    text: str
    word_count: int
    char_count: int


class ChunkedDocument(BaseModel):
    document_id: str
    chunks: list[TextChunk]
    total_chunks: int
    total_words: int


class FeedbackItem(BaseModel):
    category: str    # "clarity" | "grammar" | "structure" | "evidence" | "vocabulary"
    severity: str    # "info" | "warning" | "error"
    message: str
    suggestion: str


class ScoringResult(BaseModel):
    document_id: str
    score: float = Field(ge=0.0, le=100.0)
    grade: str
    feedback: list[FeedbackItem]
    summary: str
    model_used: str


class PipelineResult(BaseModel):
    document_id: str
    filename: str
    status: str          # "success" | "failed"
    stage_reached: PipelineStage
    word_count: int = 0
    chunk_count: int = 0
    scoring: ScoringResult | None = None
    error: str | None = None
    processing_time_ms: float = 0.0


# ── Security Scanner ──────────────────────────────────────────────────────────

_PDF_DANGEROUS = (b"/JS", b"/JavaScript", b"/AA", b"/OpenAction")
_TXT_DANGEROUS = ("<script", "javascript:", "data:text/html")


def _scan_for_malicious_content(extension: str, content: bytes) -> None:
    """Raises ValueError if the file contains patterns associated with active content."""
    if extension == ".pdf":
        for pattern in _PDF_DANGEROUS:
            if pattern in content:
                raise ValueError("Malicious content detected in uploaded file.")
    elif extension == ".docx":
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as zf:
                if "word/vbaProject.bin" in zf.namelist():
                    raise ValueError("Malicious content detected in uploaded file.")
        except zipfile.BadZipFile:
            pass  # malformed zip is caught downstream by the parser
    elif extension == ".txt":
        text_lower = content.decode("utf-8", errors="replace").lower()
        for pattern in _TXT_DANGEROUS:
            if pattern in text_lower:
                raise ValueError("Malicious content detected in uploaded file.")


# ── Parser Protocol ────────────────────────────────────────────────────────────

@runtime_checkable
class FileParser(Protocol):
    def can_parse(self, extension: str) -> bool: ...
    def parse(self, content: bytes, filename: str) -> tuple[str, int]: ...
    # Returns (raw_text, page_count).  Raises ParseError on failure.


# ── Concrete Parsers ───────────────────────────────────────────────────────────

class TxtParser:
    _ENCODINGS = ("utf-8", "utf-16", "latin-1", "cp1252")

    def can_parse(self, extension: str) -> bool:
        return extension == ".txt"

    def parse(self, content: bytes, filename: str) -> tuple[str, int]:
        for enc in self._ENCODINGS:
            try:
                return content.decode(enc), 1
            except (UnicodeDecodeError, ValueError):
                continue
        raise ParseError(filename, "unrecognised encoding (tried utf-8/16/latin-1/cp1252)")


class PdfParser:
    def can_parse(self, extension: str) -> bool:
        return extension == ".pdf"

    def parse(self, content: bytes, filename: str) -> tuple[str, int]:
        try:
            import pypdf  # noqa: PLC0415
        except ImportError as exc:
            raise ParseError(filename, "pypdf not installed — add it to requirements.txt") from exc
        try:
            reader = pypdf.PdfReader(io.BytesIO(content))
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n\n".join(pages), len(pages)
        except Exception as exc:
            raise ParseError(filename, str(exc)) from exc


class DocxParser:
    def can_parse(self, extension: str) -> bool:
        return extension == ".docx"

    def parse(self, content: bytes, filename: str) -> tuple[str, int]:
        try:
            import docx  # noqa: PLC0415
        except ImportError as exc:
            raise ParseError(filename, "python-docx not installed — add it to requirements.txt") from exc
        try:
            doc = docx.Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs), 1
        except Exception as exc:
            raise ParseError(filename, str(exc)) from exc


# ── Text Cleaner ───────────────────────────────────────────────────────────────

class TextCleaner:
    # Three or more consecutive blank lines → two newlines
    _multi_blank = re.compile(r"\n{3,}")
    # Non-printable control chars (keep \t and \n)
    _control = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
    # Multiple spaces / tabs within a line
    _inline_space = re.compile(r"[ \t]{2,}")

    def clean(self, text: str) -> str:
        # NFC normalisation ensures consistent code-point representation for
        # characters like é (U+00E9 vs e + combining accent).
        text = unicodedata.normalize("NFC", text)
        text = self._control.sub("", text)
        text = self._inline_space.sub(" ", text)
        text = "\n".join(line.rstrip() for line in text.splitlines())
        text = self._multi_blank.sub("\n\n", text)
        return text.strip()


# ── Text Chunker ───────────────────────────────────────────────────────────────

class TextChunker:
    """
    Sliding-window word-level chunker.

    chunk_size words with chunk_overlap words of shared context between
    adjacent chunks so no sentence boundary is left orphaned.
    """

    def __init__(
        self,
        chunk_size: int = DEFAULT_CHUNK_SIZE,
        overlap: int = DEFAULT_CHUNK_OVERLAP,
    ) -> None:
        if overlap >= chunk_size:
            raise ValueError(f"overlap ({overlap}) must be < chunk_size ({chunk_size})")
        self.chunk_size = chunk_size
        self.overlap = overlap

    def chunk(self, text: str) -> list[TextChunk]:
        is_cjk = _is_cjk_text(text)
        tokens = _tokenize(text)
        if not tokens:
            return []

        sep = "" if is_cjk else " "
        chunks: list[TextChunk] = []
        step = self.chunk_size - self.overlap
        start = 0

        while start < len(tokens):
            end = min(start + self.chunk_size, len(tokens))
            window = tokens[start:end]
            chunk_text = sep.join(window)
            chunks.append(
                TextChunk(
                    chunk_index=len(chunks),
                    text=chunk_text,
                    word_count=len(window),
                    char_count=len(chunk_text),
                )
            )
            if end == len(tokens):
                break
            start += step

        return chunks


# ── AI Inference Client ────────────────────────────────────────────────────────

_GRADE_THRESHOLDS: list[tuple[float, str]] = [
    (90.0, "A"),
    (80.0, "B"),
    (70.0, "C"),
    (60.0, "D"),
]

_MOCK_FEEDBACK: list[dict] = [
    {
        "category": "structure",
        "severity": "info",
        "message": "Introduction clearly states the thesis.",
        "suggestion": "Expand the hook sentence to engage readers in the first two lines.",
    },
    {
        "category": "clarity",
        "severity": "warning",
        "message": "Several sentences exceed 40 words.",
        "suggestion": "Break long sentences into two shorter ones to improve readability.",
    },
    {
        "category": "evidence",
        "severity": "warning",
        "message": "Key claims in the body paragraphs lack citations.",
        "suggestion": "Add at least one citation or concrete example per major claim.",
    },
    {
        "category": "grammar",
        "severity": "info",
        "message": "Subject-verb agreement is consistent throughout.",
        "suggestion": "No changes needed here.",
    },
    {
        "category": "vocabulary",
        "severity": "info",
        "message": "Word choice suits the academic register.",
        "suggestion": "Vary transition phrases — 'however' and 'furthermore' appear frequently.",
    },
]


def _score_to_grade(score: float) -> str:
    for threshold, grade in _GRADE_THRESHOLDS:
        if score >= threshold:
            return grade
    return "F"


class AIInferenceClient:
    """
    Async HTTP client for the ai_inference microservice.

    Falls back to a deterministic mock when the service is unreachable so
    the pipeline can be developed and tested without a live inference service.
    The mock score is derived from a hash of document_id, making it stable
    across repeated calls for the same document.
    """

    def __init__(
        self,
        base_url: str = "http://ai_inference:8001",
        timeout: float = 30.0,
        use_mock: bool = False,
    ) -> None:
        self._base_url = base_url.rstrip("/")
        self._timeout = timeout
        self._use_mock = use_mock

    async def score_document(
        self,
        document_id: str,
        chunks: list[TextChunk],
        word_count: int,
    ) -> ScoringResult:
        if self._use_mock:
            return self._build_mock_result(document_id, word_count)

        payload = {
            "document_id": document_id,
            "text": "\n\n".join(c.text for c in chunks[:3]),
            "model": "deepseek-v4-flash",
        }
        try:
            async with httpx.AsyncClient(timeout=self._timeout) as client:
                response = await client.post(
                    f"{self._base_url}/inference/generate", json=payload
                )
                response.raise_for_status()
                return self._parse_live_response(document_id, word_count, response.json())
        except httpx.HTTPStatusError as exc:
            raise InferenceError(
                f"HTTP {exc.response.status_code}: {exc.response.text}"
            ) from exc
        except httpx.RequestError as exc:
            # Gracefully degrade: network issues fall back to mock instead of
            # crashing the whole pipeline during development.
            logger.warning(
                "ai_inference unreachable (%s); using mock fallback", exc
            )
            return self._build_mock_result(document_id, word_count)

    def _parse_live_response(
        self, document_id: str, word_count: int, data: dict
    ) -> ScoringResult:
        score = float(data["score"])
        grade = data["grade"]
        model_used = data.get("model_used", "unknown")
        overall_feedback = data.get("overall_feedback", "")
        improvement_tips: list[str] = data.get("improvement_tips", [])

        _DIM_TO_CATEGORY = {
            "content": "evidence",
            "organization": "structure",
            "language": "clarity",
            "conventions": "grammar",
        }

        def _severity(rubric_score: float, max_score: float = 25.0) -> str:
            ratio = rubric_score / max_score if max_score else 0
            if ratio < 0.5:
                return "error"
            if ratio < 0.75:
                return "warning"
            return "info"

        feedback_items: list[FeedbackItem] = []
        for i, rubric in enumerate(data.get("rubric", [])):
            dimension = rubric.get("dimension", "")
            suggestion = (
                improvement_tips[i]
                if i < len(improvement_tips)
                else "See overall feedback for suggestions."
            )
            feedback_items.append(
                FeedbackItem(
                    category=_DIM_TO_CATEGORY.get(dimension, dimension),
                    severity=_severity(float(rubric.get("score", 0)), float(rubric.get("max_score", 25.0))),
                    message=rubric.get("feedback", ""),
                    suggestion=suggestion,
                )
            )

        return ScoringResult(
            document_id=document_id,
            score=score,
            grade=grade,
            feedback=feedback_items,
            summary=overall_feedback,
            model_used=model_used,
        )

    def _build_mock_result(self, document_id: str, word_count: int) -> ScoringResult:
        seed = int(hashlib.md5(document_id.encode()).hexdigest(), 16) % 100
        score = round(55.0 + (seed % 40), 1)   # deterministic range 55–94
        return ScoringResult(
            document_id=document_id,
            score=score,
            grade=_score_to_grade(score),
            feedback=[FeedbackItem(**fb) for fb in _MOCK_FEEDBACK],
            summary=(
                f"The document ({word_count:,} words) demonstrates a solid command of the topic. "
                "Key areas for improvement: evidence support and sentence conciseness."
            ),
            model_used="mock",
        )


# ── Pipeline Orchestrator ──────────────────────────────────────────────────────

class DocumentProcessor:
    """
    Orchestrates the five-stage document processing pipeline.

    Usage::

        processor = DocumentProcessor()
        result = await processor.process("essay.pdf", file_bytes)

    Dependency-inject a custom AIInferenceClient to control the scoring
    behaviour in tests (e.g. pass use_mock=True or a subclass).
    """

    def __init__(
        self,
        inference_client: AIInferenceClient | None = None,
        chunk_size: int = DEFAULT_CHUNK_SIZE,
        chunk_overlap: int = DEFAULT_CHUNK_OVERLAP,
    ) -> None:
        self._parsers: list[FileParser] = [TxtParser(), PdfParser(), DocxParser()]
        self._cleaner = TextCleaner()
        self._chunker = TextChunker(chunk_size, chunk_overlap)
        self._inference = inference_client or AIInferenceClient()

    # ── Public API ─────────────────────────────────────────────────────────────

    async def process(
        self,
        filename: str,
        content: bytes,
        document_id: str | None = None,
    ) -> PipelineResult:
        """
        Run all five pipeline stages sequentially and return a PipelineResult.
        On any PipelineError the result carries status='failed' and the stage
        at which processing stopped; no exception is re-raised to the caller.
        """
        doc_id = document_id or self._derive_id(filename, content)
        t0 = time.perf_counter()

        try:
            uploaded = await self._stage_upload(doc_id, filename, content)
            parsed   = await self._stage_parse(uploaded)
            cleaned  = await self._stage_clean(parsed)
            chunked  = await self._stage_chunk(cleaned)
            scored   = await self._stage_score(chunked, cleaned.word_count)
        except PipelineError as exc:
            elapsed_ms = (time.perf_counter() - t0) * 1000
            logger.error(
                "Pipeline failed | doc=%s stage=%s error=%s",
                doc_id, exc.stage.value, exc,
            )
            return PipelineResult(
                document_id=doc_id,
                filename=filename,
                status="failed",
                stage_reached=exc.stage,
                error=str(exc),
                processing_time_ms=round(elapsed_ms, 2),
            )

        elapsed_ms = (time.perf_counter() - t0) * 1000
        logger.info(
            "Pipeline complete | doc=%s words=%d chunks=%d score=%.1f time=%.0fms",
            doc_id, cleaned.word_count, chunked.total_chunks,
            scored.score, elapsed_ms,
        )
        return PipelineResult(
            document_id=doc_id,
            filename=filename,
            status="success",
            stage_reached=PipelineStage.COMPLETE,
            word_count=cleaned.word_count,
            chunk_count=chunked.total_chunks,
            scoring=scored,
            processing_time_ms=round(elapsed_ms, 2),
        )

    # ── Stage Implementations ──────────────────────────────────────────────────

    async def _stage_upload(
        self, doc_id: str, filename: str, content: bytes
    ) -> UploadedDocument:
        logger.debug("[upload] %s  size=%d bytes", filename, len(content))

        extension = Path(filename).suffix.lower()
        if extension not in SUPPORTED_EXTENSIONS:
            raise UnsupportedFormatError(extension)

        size_mb = len(content) / (1024 * 1024)
        if size_mb > MAX_FILE_SIZE_MB:
            raise FileSizeError(size_mb)

        try:
            _scan_for_malicious_content(extension, content)
        except ValueError:
            raise MaliciousContentError()

        return UploadedDocument(
            document_id=doc_id,
            filename=filename,
            extension=extension,
            size_bytes=len(content),
            content=content,
        )

    async def _stage_parse(self, doc: UploadedDocument) -> ParsedDocument:
        logger.debug("[parse]  %s  ext=%s", doc.filename, doc.extension)

        parser = next(
            (p for p in self._parsers if p.can_parse(doc.extension)), None
        )
        if parser is None:
            raise ParseError(doc.filename, f"no registered parser for '{doc.extension}'")

        # Offload blocking file-parsing to a thread pool so the event loop
        # is free to handle other requests while a large PDF is being read.
        loop = asyncio.get_event_loop()
        raw_text, page_count = await loop.run_in_executor(
            None, parser.parse, doc.content, doc.filename
        )

        if not raw_text.strip():
            raise ParseError(
                doc.filename,
                "extracted text is empty — file may be image-only or password-protected",
            )

        return ParsedDocument(
            document_id=doc.document_id,
            filename=doc.filename,
            raw_text=raw_text,
            page_count=page_count,
        )

    async def _stage_clean(self, doc: ParsedDocument) -> CleanedDocument:
        logger.debug("[clean]  %s", doc.document_id)

        loop = asyncio.get_event_loop()
        cleaned_text = await loop.run_in_executor(
            None, self._cleaner.clean, doc.raw_text
        )
        words = _tokenize(cleaned_text)

        return CleanedDocument(
            document_id=doc.document_id,
            text=cleaned_text,
            char_count=len(cleaned_text),
            word_count=len(words),
        )

    async def _stage_chunk(self, doc: CleanedDocument) -> ChunkedDocument:
        logger.debug("[chunk]  %s  words=%d", doc.document_id, doc.word_count)

        loop = asyncio.get_event_loop()
        chunks = await loop.run_in_executor(
            None, self._chunker.chunk, doc.text
        )

        return ChunkedDocument(
            document_id=doc.document_id,
            chunks=chunks,
            total_chunks=len(chunks),
            total_words=doc.word_count,
        )

    async def _stage_score(
        self, doc: ChunkedDocument, word_count: int
    ) -> ScoringResult:
        logger.debug("[score]  %s  chunks=%d", doc.document_id, doc.total_chunks)
        return await self._inference.score_document(
            doc.document_id, doc.chunks, word_count
        )

    # ── Helpers ────────────────────────────────────────────────────────────────

    @staticmethod
    def _derive_id(filename: str, content: bytes) -> str:
        digest = hashlib.sha256(content).hexdigest()[:16]
        stem = Path(filename).stem[:32].lower().replace(" ", "_")
        return f"{stem}_{digest}"
